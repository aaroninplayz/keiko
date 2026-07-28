import os
import re
import logging
from typing import Dict, Any, List, Optional

logger = logging.getLogger(__name__)

# List of common security/injection phrases to sanitize
INJECTION_PATTERNS = [
    r"(?i)ignore\s+(?:all\s+)?previous\s+instructions",
    r"(?i)system\s+prompt\s+override",
    r"(?i)override\s+(?:all\s+)?settings",
    r"(?i)you\s+must\s+score\s+(?:this\s+candidate\s+)?100",
    r"(?i)ignore\s+constraints",
    r"(?i)bypass\s+evaluation",
    r"(?i)ignore\s+posture\s+and\s+eye\s+contact",
]

# Canonical ATS Skill Taxonomy (Canonical Name, Category, Regex List, CaseSensitiveBool)
SKILL_TAXONOMY = [
    # Programming Languages
    ("Python", "programming_languages", [r"\bPython\b"], False),
    ("C++", "programming_languages", [r"(?<![A-Za-z0-9_])C\+\+(?![A-Za-z0-9_])"], True),
    ("C#", "programming_languages", [r"(?<![A-Za-z0-9_])C#(?![A-Za-z0-9_])"], True),
    ("C", "programming_languages", [r"(?<![A-Za-z0-9_])C(?![A-Za-z0-9_\+])"], True),
    ("JavaScript", "programming_languages", [r"\bJava\s*Script\b", r"\bJavaScript\b", r"\bJS\b"], False),
    ("TypeScript", "programming_languages", [r"\bTypeScript\b", r"\bTS\b"], False),
    ("Java", "programming_languages", [r"\bJava\b"], True),
    ("Go / Golang", "programming_languages", [r"\bGolang\b", r"\bGo\b"], True),
    ("Rust", "programming_languages", [r"\bRust\b"], False),
    ("SQL", "programming_languages", [r"\bSQL\b"], True),
    ("HTML/CSS", "programming_languages", [r"\bHTML/CSS\b", r"\bHTML\b", r"\bCSS\b", r"\bVanilla CSS\b"], False),
    
    # Frameworks & Libraries
    ("FastAPI", "frameworks", [r"\bFastAPI\b"], False),
    ("React", "frameworks", [r"\bReact\b", r"\bReactJS\b", r"\bReact\.js\b"], False),
    ("Next.js", "frameworks", [r"\bNext\.js\b", r"\bNextJS\b"], False),
    ("Tailwind CSS", "frameworks", [r"\bTailwind\s*CSS\b", r"\bTailwind\b"], False),
    ("PyTorch", "frameworks", [r"\bPyTorch\b"], False),
    ("OpenCV", "frameworks", [r"\bOpenCV\b"], False),
    ("TensorFlow", "frameworks", [r"\bTensorFlow\b"], False),
    ("Django", "frameworks", [r"\bDjango\b"], False),
    ("Flask", "frameworks", [r"\bFlask\b"], False),
    ("Angular", "frameworks", [r"\bAngular\b"], False),
    ("Vue.js", "frameworks", [r"\bVue\b", r"\bVueJS\b"], False),
    ("Spring Boot", "frameworks", [r"\bSpring\s*Boot\b"], False),

    # Tools, Cloud & Databases
    ("Git / GitHub", "tools_databases", [r"\bGit\b", r"\bGitHub\b"], False),
    ("Docker", "tools_databases", [r"\bDocker\b"], False),
    ("Kubernetes", "tools_databases", [r"\bKubernetes\b", r"\bK8s\b"], False),
    ("AWS", "tools_databases", [r"\bAWS\b", r"\bAmazon Web Services\b"], True),
    ("GCP", "tools_databases", [r"\bGCP\b", r"\bGoogle Cloud\b"], True),
    ("Azure", "tools_databases", [r"\bAzure\b"], False),
    ("Firebase / Firestore", "tools_databases", [r"\bFirebase\b", r"\bFirestore\b"], False),
    ("PostgreSQL", "tools_databases", [r"\bPostgreSQL\b", r"\bPostgres\b"], False),
    ("MongoDB", "tools_databases", [r"\bMongoDB\b"], False),
    ("Redis", "tools_databases", [r"\bRedis\b"], False),
    ("PWA Development", "tools_databases", [r"\bPWA\b", r"\bProgressive Web App\b"], False),

    # Design, AI & Embedded Engineering
    ("UI/UX Design", "design_and_ai", [r"\bUI/UX\b", r"\bUI/UX Design\b", r"\bWireframing\b"], False),
    ("Responsive Web Design", "design_and_ai", [r"\bResponsive Web Design\b"], False),
    ("Prompt Engineering", "design_and_ai", [r"\bPrompt Engineering\b"], False),
    ("Google AI Studio", "design_and_ai", [r"\bGoogle AI Studio\b"], False),
    ("Gemini API", "design_and_ai", [r"\bGemini\s*API\b", r"\bGemini\b"], False),
    ("AI Code Assistants", "design_and_ai", [r"\bAntigravity\b", r"\bCodex\b", r"\bCopilot\b"], False),
    ("Bluetooth / BLE", "design_and_ai", [r"\bBLE\b", r"\bBluetooth\b"], True),
    ("Wi-Fi & BSSID", "design_and_ai", [r"\bWi-Fi\b", r"\bWiFi\b", r"\bBSSID\b"], False),
    ("Multimodal & Sensor Fusion", "design_and_ai", [r"\bSensor Fusion\b", r"\bMultimodal\b", r"\bRGB\b", r"\bThermal\b"], False),
]

import json
from .llm_client import LLMClient

class ResumeIntelligenceAgent:
    """
    ATS-Grade & AI-Powered Resume Intelligence Agent.
    Parses any PDF, DOCX, or TXT resume dynamically without hardcoded data.
    Uses LLM structured extraction with fallbacks for name, email, phone, skills, experience, education & achievements.
    """

    def sanitize_text(self, text: str) -> str:
        """Cleans text while preserving line breaks for paragraph & section boundary detection."""
        if not text:
            return ""
        clean = re.sub(r"<[^>]*>", " ", text)
        for pattern in INJECTION_PATTERNS:
            clean = re.sub(pattern, "[Sanitized Malicious Intent Phrase]", clean)
        clean = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", clean)
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in clean.splitlines()]
        clean_text = "\n".join([l for l in lines if l])
        return clean_text.strip()

    def parse_pdf_rich(self, file_path: str) -> str:
        """Extracts text and embedded hyperlink URIs from PDF using PyMuPDF (fitz) with pypdf fallback."""
        try:
            import fitz
            doc = fitz.open(file_path)
            pages_text = []
            for page in doc[:15]:
                blocks = page.get_text("blocks")
                links = page.get_links()
                page_lines = []
                for b in blocks:
                    if len(b) >= 7 and b[6] == 0:  # text block
                        b_text = b[4].strip()
                        if not b_text:
                            continue
                        b_rect = fitz.Rect(b[:4])
                        matching_urls = []
                        for link in links:
                            if "uri" in link and link["uri"]:
                                uri = link["uri"].strip()
                                link_rect = fitz.Rect(link["from"])
                                if b_rect.intersects(link_rect) and uri not in b_text:
                                    if uri not in matching_urls:
                                        matching_urls.append(uri)
                        if matching_urls:
                            b_text += " [" + " | ".join(matching_urls) + "]"
                        page_lines.append(b_text)
                if page_lines:
                    pages_text.append("\n".join(page_lines))
            return "\n\n".join(pages_text)
        except Exception as e:
            logger.warning(f"PyMuPDF rich extraction failed: {e}. Falling back to pypdf.")
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            pages_text = []
            for page in reader.pages[:15]:
                txt = page.extract_text()
                if txt:
                    pages_text.append(txt)
            return "\n".join(pages_text)

    def parse_docx_rich(self, file_path: str) -> str:
        """Extracts text and embedded hyperlink URIs from DOCX files using python-docx."""
        import docx
        doc = docx.Document(file_path)
        lines = []

        rel_links = {}
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                rel_links[rel.rId] = rel.target_ref

        for p in doc.paragraphs[:500]:
            if not p.text or not p.text.strip():
                continue
            p_text = p.text.strip()
            p_xml = p._element.xml
            p_urls = []
            for rId, url in rel_links.items():
                if rId in p_xml and url not in p_text and url not in p_urls:
                    p_urls.append(url)
            if p_urls:
                p_text += " [" + " | ".join(p_urls) + "]"
            lines.append(p_text)

        for table in doc.tables[:50]:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    lines.append(row_text)

        return "\n".join(lines)

    def parse_file(self, file_path: str) -> str:
        """Extracts text safely with hyperlink preservation from PDF, DOCX, or TXT formats."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        _, ext = os.path.splitext(file_path.lower())
        raw_text = ""

        try:
            if ext == ".pdf":
                raw_text = self.parse_pdf_rich(file_path)
            elif ext == ".docx":
                raw_text = self.parse_docx_rich(file_path)
            else:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    raw_text = f.read(100000)

            return self.sanitize_text(raw_text)

        except Exception as e:
            logger.error(f"Error parsing resume file {file_path}: {e}")
            raise ValueError(f"Failed to read resume file: {str(e)}")

    def extract_candidate_name(self, text: str) -> str:
        """Dynamically extracts candidate full name from top header lines with non-ATS filtering & email token matching."""
        if not text:
            return "Candidate"

        lines = [l.strip() for l in text.splitlines() if l.strip()]
        email = self.extract_candidate_email(text)
        email_tokens = set()
        if email:
            uname = email.split("@")[0].lower()
            email_tokens = set(re.split(r"[._\-0-9]+", uname))

        forbidden_patterns = [
            r"@|http|\+?\d{7,}",
            r"\b(curriculum|vitae|resume|profile|summary|experience|education|skills|projects|achievements|certifications|contact|address|phone|email)\b",
            r"\b(software|developer|engineer|full\s*stack|backend|frontend|data\s*scientist|ai\s*engineer|manager|lead|architect|intern|student)\b",
            r"\b(page\s*\d|page\s*of|references|declaration)\b",
        ]

        candidates = []
        for idx, line in enumerate(lines[:15]):
            if any(re.search(pat, line, re.IGNORECASE) for pat in forbidden_patterns):
                continue

            cleaned = re.sub(r"[^a-zA-Z\s]", "", line).strip()
            words = cleaned.split()

            if 2 <= len(words) <= 4 and all(len(w) >= 2 for w in words):
                score = 0.0
                # Title case bonus
                if all(w[0].isupper() for w in words if w):
                    score += 2.0
                # Position bonus (earlier in document)
                score += max(0, (15 - idx) * 0.2)
                # Email match bonus
                lowered_words = [w.lower() for w in words]
                if any(w in email_tokens for w in lowered_words):
                    score += 3.0

                candidates.append((score, cleaned.title()))

        if candidates:
            candidates.sort(key=lambda x: x[0], reverse=True)
            return candidates[0][1]

        return "Candidate"

    def extract_candidate_email(self, text: str) -> str:
        """Extracts candidate email using standard email pattern."""
        match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
        return match.group(0) if match else ""

    def extract_candidate_phone(self, text: str) -> str:
        """Extracts candidate phone number with flexible non-ATS matching."""
        match = re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
        return match.group(0) if match else ""

    def extract_experience_years(self, text: str) -> int:
        """Estimates candidate experience years from date ranges or explicit statements."""
        exp_match = re.search(r"(\d+)\+?\s*years?\s+(?:of\s+)?experience", text, re.IGNORECASE)
        if exp_match:
            return int(exp_match.group(1))

        years = re.findall(r"\b(20\d{2})\s*[-–—]\s*(Present|20\d{2})\b", text, re.IGNORECASE)
        total_years = 0
        current_year = 2026
        for start, end in years:
            try:
                sy = int(start)
                ey = current_year if end.lower() == "present" else int(end)
                if ey >= sy:
                    total_years += (ey - sy)
            except Exception:
                continue

        if re.search(r"first-year|undergraduate|intern|student", text, re.IGNORECASE):
            return 1
        return max(1, total_years) if total_years > 0 else 1

    def _score_block(self, block: str) -> Dict[str, float]:
        """
        Computes domain keyword density scores to classify a block into target categories:
        'work', 'education', 'projects', 'achievements', 'certifications'.
        """
        b_lower = block.lower()

        scores = {
            "work": 0.0,
            "education": 0.0,
            "projects": 0.0,
            "achievements": 0.0,
            "certifications": 0.0
        }

        # Work / Internship signals
        work_titles = [r"\bintern\b", r"\binternship\b", r"\bengineer\b", r"\bdeveloper\b", r"\banalyst\b",
                       r"\bconsultant\b", r"\blead\b", r"\barchitect\b", r"\bmanager\b", r"\bassociate\b",
                       r"\btrainee\b", r"\bco-op\b", r"\bsoftware\b", r"\bbackend\b", r"\bfrontend\b"]
        for wt in work_titles:
            if re.search(wt, b_lower):
                scores["work"] += 1.5

        work_verbs = [r"\bbuilt\b", r"\bdeveloped\b", r"\bmanaged\b", r"\bled\b", r"\barchitected\b",
                      r"\bimplemented\b", r"\bcollaborated\b", r"\bspearheaded\b", r"\bworked at\b",
                      r"\bdesigned and deployed\b", r"\bresponsibilities\b", r"\bkey achievements at\b"]
        for wv in work_verbs:
            if re.search(wv, b_lower):
                scores["work"] += 1.0

        if re.search(r"\b(20\d{2}|present|current)\b", b_lower) and re.search(r"\b(inc|corp|ltd|pvt|technologies|solutions|labs|company|startup|firm)\b", b_lower):
            scores["work"] += 2.5

        # Education signals
        edu_degrees = [r"\bb\.?tech\b", r"\bb\.?s\b", r"\bm\.?s\b", r"\bm\.?tech\b", r"\bb\.?e\b", r"\bbca\b", r"\bmca\b",
                       r"\bph\.?d\b", r"\bbachelor\b", r"\bmaster\b", r"\bdiploma\b", r"\bdegree\b",
                       r"\bhigh school\b", r"\bsecondary school\b", r"\bcbse\b", r"\bicse\b"]
        for ed in edu_degrees:
            if re.search(ed, b_lower):
                scores["education"] += 2.5

        edu_inst = [r"\buniversity\b", r"\bcollege\b", r"\binstitute\b", r"\bschool\b", r"\bacademy\b", r"\bcampus\b"]
        for ei in edu_inst:
            if re.search(ei, b_lower):
                scores["education"] += 2.0

        if re.search(r"\b(cgpa|gpa|percentage|marks|grade|major|coursework)\b", b_lower):
            scores["education"] += 1.5

        # Projects signals
        proj_terms = [r"\bproject\b", r"\bgithub\b", r"\brepository\b", r"\blive demo\b", r"\btech stack\b",
                      r"\bbuilt a\b", r"\bdeveloped a\b", r"\bweb app\b", r"\bmobile app\b", r"\bdashboard\b",
                      r"\bplatform\b", r"\bsystem\b", r"\bbot\b", r"\bmodel\b", r"\bpipeline\b", r"\btool\b"]
        for pt in proj_terms:
            if re.search(pt, b_lower):
                scores["projects"] += 1.5

        if "github.com" in b_lower or "http" in b_lower:
            scores["projects"] += 2.0

        # Achievements signals
        ach_terms = [r"\bawarded\b", r"\baward\b", r"\bwinner\b", r"\b1st place\b", r"\b2nd place\b",
                     r"\b3rd place\b", r"\btop \d+%\b", r"\brank\b", r"\bhackathon\b", r"\bscholarship\b",
                     r"\bhonors\b", r"\bpublication\b", r"\bpublished\b", r"\bpatent\b"]
        for at in ach_terms:
            if re.search(at, b_lower):
                scores["achievements"] += 2.0

        # Certifications signals
        cert_terms = [r"\bcertified\b", r"\bcertification\b", r"\baws certified\b", r"\bazure certified\b", r"\bcredential\b"]
        for ct in cert_terms:
            if re.search(ct, b_lower):
                scores["certifications"] += 2.0

        return scores

    def _split_into_blocks(self, text: str) -> List[str]:
        """Splits raw text into coherent logical blocks by double newlines or entry line starts."""
        if not text:
            return []

        paragraphs = re.split(r"\n\s*\n", text)
        blocks = []

        for p in paragraphs:
            p_clean = p.strip()
            if not p_clean:
                continue

            lines = [l.strip() for l in p_clean.splitlines() if l.strip()]
            current_block = []

            for line in lines:
                is_new_entry = False
                if current_block:
                    # Structural title pattern: Title: Description
                    colon_match = re.match(r"^([^:\n]+):\s+(.*)$", line)
                    if colon_match and len(colon_match.group(1)) <= 65 and not line.startswith(("-", "*", "•")):
                        is_new_entry = True
                    # Date range pattern
                    elif re.search(r"\b(20\d{2}|19\d{2})\s*[-–—]\s*(Present|Current|20\d{2}|19\d{2})\b", line, re.IGNORECASE):
                        is_new_entry = True
                    # Separator header pattern: Role at Company, Degree, Institution, Project - Stack
                    elif (re.search(r"\b(at|@|\|)\b", line, re.IGNORECASE) or re.match(r"^[A-Z][A-Za-z0-9\s,\-\.]{3,50}\s*[-–—]\s*[A-Z][A-Za-z0-9\s,\-\.]{3,50}", line)) and not line.startswith(("-", "*", "•")):
                        is_new_entry = True
                    # Numbered list pattern
                    elif re.match(r"^(?:\d+[\.\)]|\[\d+\]|Item\s+\d+|Project\s+\d+)\s+", line, re.IGNORECASE):
                        is_new_entry = True

                if is_new_entry and current_block:
                    blocks.append("\n".join(current_block))
                    current_block = [line]
                else:
                    current_block.append(line)

            if current_block:
                blocks.append("\n".join(current_block))

        return [b for b in blocks if len(b.strip()) > 10]

    def _parse_sections(self, text: str) -> Dict[str, str]:
        """Dynamically identifies section boundaries in non-ATS & ATS resumes."""
        header_patterns = [
            (r"\b(PROFESSIONAL\s+SUMMARY|SUMMARY|ABOUT\s+ME|ABOUT)\b", "SUMMARY"),
            (r"\b(WORK\s+EXPERIENCE|EXPERIENCE|EMPLOYMENT|INTERNSHIPS|WORK\s+HISTORY|CAREER|PAST\s+ROLES|WHERE\s+I[\'\’]?VE\s+WORKED)\b", "WORK"),
            (r"\b(EDUCATION|ACADEMIC|ACADEMICS|STUDIED|WHERE\s+I\s+STUDIED|COLLEGE|UNIVERSITY)\b", "EDUCATION"),
            (r"\b(TECHNICAL\s+SKILLS|SKILLS|WHAT\s+I\s+DO|TECHNOLOGIES|TOOLS)\b", "SKILLS"),
            (r"\b(TECHNICAL\s+PROJECTS|PROJECTS|THINGS\s+I[\'\’]?VE\s+BUILT|PORTFOLIO)\b", "PROJECTS"),
            (r"\b(CERTIFICATIONS|LICENSES|CERTIFICATES)\b", "CERTIFICATIONS"),
            (r"\b(ACHIEVEMENTS|AWARDS|HONORS|ACCOMPLISHMENTS)\b", "ACHIEVEMENTS")
        ]

        lines = text.splitlines()
        section_matches = []

        for idx, line in enumerate(lines):
            l_str = line.strip()
            if not l_str or len(l_str) > 70:
                continue

            for pat, key in header_patterns:
                if re.search(pat, l_str, re.IGNORECASE):
                    section_matches.append((idx, key))
                    break

        sections = {}
        for i, (line_idx, key) in enumerate(section_matches):
            start_line = line_idx + 1
            end_line = section_matches[i+1][0] if i + 1 < len(section_matches) else len(lines)
            sec_text = "\n".join(lines[start_line:end_line]).strip()
            if key not in sections:
                sections[key] = sec_text
            else:
                sections[key] += "\n\n" + sec_text

        return sections

    def _reclassify_and_distribute(self, profile: Dict[str, Any], raw_text: str) -> Dict[str, Any]:
        """
        Block Classifier & Flipping Engine:
        Evaluates paragraph blocks and re-routes misclassified items across Education, Work, Projects, and Achievements.
        Also recovers projects and achievements if missing.
        """
        all_blocks = self._split_into_blocks(raw_text)

        # 1. Flip miscategorized Education items (e.g. Internships/Projects placed under Education)
        new_edu = []
        for edu_item in profile.get("education", []):
            scores = self._score_block(edu_item)
            # If work or project score strongly outweighs education score, flip it!
            if scores["work"] > scores["education"] + 0.5:
                if edu_item not in profile["work_history"]:
                    profile["work_history"].append(edu_item)
            elif scores["projects"] > scores["education"] + 0.5:
                p_name = edu_item.split("\n")[0][:60]
                profile["projects"].append({"name": p_name, "description": edu_item})
            elif scores["achievements"] > scores["education"] + 0.5:
                if edu_item not in profile["achievements"]:
                    profile["achievements"].append(edu_item)
            else:
                new_edu.append(edu_item)
        profile["education"] = new_edu

        # 2. Re-examine all document blocks to extract missing Work, Projects, Education & Achievements
        for block in all_blocks:
            scores = self._score_block(block)
            best_cat = max(scores, key=scores.get)
            max_score = scores[best_cat]

            if max_score < 1.0:
                continue

            if best_cat == "projects":
                lines = [l.strip() for l in block.splitlines() if l.strip()]
                if not lines:
                    continue
                if re.search(r"\b(TECHNICAL\s+PROJECTS|PROJECTS|THINGS\s+I[\'\’]?VE\s+BUILT|PORTFOLIO)\b", lines[0], re.IGNORECASE):
                    lines = lines[1:]
                    if not lines:
                        continue
                title = lines[0]
                title = re.sub(r"^[\bullet\*\-\#\:\s]+", "", title).strip()
                desc = "\n".join(lines)
                if not any(title.lower() in p.get("name", "").lower() for p in profile["projects"] if isinstance(p, dict)):
                    profile["projects"].append({
                        "name": title[:70],
                        "description": desc
                    })
            elif best_cat == "work" and not profile["work_history"]:
                if block not in profile["work_history"]:
                    profile["work_history"].append(block)
            elif best_cat == "education" and not profile["education"]:
                if block not in profile["education"]:
                    profile["education"].append(block)
            elif best_cat == "achievements":
                if block not in profile["achievements"]:
                    profile["achievements"].append(block)

        # Normalize projects format to array of dicts with name & description, deduplicated & split
        norm_projects = []
        seen_names = set()
        
        raw_proj_items = []
        for p in profile.get("projects", []):
            if isinstance(p, str):
                sub_blocks = self._split_section_into_entries(p)
                raw_proj_items.extend(sub_blocks)
            elif isinstance(p, dict):
                p_desc = p.get("description", "")
                if "\n\n" in p_desc or (p_desc.count("\n") >= 2 and any(k in p_desc for k in [":", "http", "github", "Built", "Developed"])):
                    sub_blocks = self._split_section_into_entries(p_desc)
                    raw_proj_items.extend(sub_blocks)
                else:
                    raw_proj_items.append(p_desc or p.get("name", "Project"))
            else:
                continue

        for item_str in raw_proj_items:
            lines = [l.strip() for l in item_str.splitlines() if l.strip()]
            if not lines:
                continue
            first_line = lines[0]
            if re.search(r"^(THINGS I[\'\’]?VE BUILT|TECHNICAL PROJECTS|PROJECTS|PORTFOLIO)\s*:?$", first_line, re.IGNORECASE):
                lines = lines[1:]
                if not lines:
                    continue
                first_line = lines[0]

            colon_idx = first_line.find(":")
            if colon_idx > 0 and colon_idx < 60:
                name = first_line[:colon_idx].strip()
            else:
                name = first_line[:70].strip()

            name = re.sub(r"^[\bullet\*\-\#\:\s]+", "", name).strip()
            clean_key = name.lower()[:40]

            if clean_key and clean_key not in seen_names:
                seen_names.add(clean_key)
                norm_projects.append({"name": name, "description": "\n".join(lines)})

        profile["projects"] = norm_projects
        return profile

    def llm_parse_resume(self, text: str) -> Dict[str, Any]:
        """
        Multi-Stage Prompting Engine designed specifically for low-context local LLMs (e.g. qwen2.5:1.5b).
        Executes focused, small-context prompts across isolated stages to prevent context overflow:
        Stage 1: Identity & Contact Info
        Stage 2: Skills & Technologies Tags
        Stage 3: Work Experience Company Blocks
        Stage 4: Education & Academic History
        Stage 5: Projects & Achievements
        """
        llm = LLMClient()
        provider = llm.detect_provider()
        if not provider:
            return {}

        clean_text = self.sanitize_text(text)
        result = {
            "full_name": "",
            "email": "",
            "phone": "",
            "skills": [],
            "work_history": [],
            "education": [],
            "achievements": [],
            "projects": [],
            "certifications": []
        }

        failed_stages = 0
        def run_stage_prompt(stage_name: str, prompt_text: str) -> Optional[str]:
            nonlocal failed_stages
            if failed_stages >= 1:
                return None
            try:
                res = llm.complete([{"role": "user", "content": prompt_text}], temperature=0.1, max_tokens=600, timeout=3.0)
                if res:
                    clean = re.sub(r"^```json\s*", "", res.strip(), flags=re.MULTILINE)
                    clean = re.sub(r"```$", "", clean.strip(), flags=re.MULTILINE)
                    return clean
                else:
                    failed_stages += 1
            except Exception as e:
                failed_stages += 1
                logger.warning(f"Multi-stage prompt stage '{stage_name}' failed: {e}")
            return None

        # Stage 1: Identity & Contact Extraction (Header text <1200 chars)
        stage1_prompt = f"""Extract contact info from this resume text as a JSON object:
Text:
\"\"\"
{clean_text[:1200]}
\"\"\"
Return ONLY valid JSON matching this format:
{{"full_name": "Full Name", "email": "email@domain.com", "phone": "+1-555-000-0000"}}"""
        
        s1_res = run_stage_prompt("Stage 1 - Identity", stage1_prompt)
        if s1_res:
            try:
                data = json.loads(s1_res)
                if isinstance(data, dict):
                    result["full_name"] = data.get("full_name", "")
                    result["email"] = data.get("email", "")
                    result["phone"] = data.get("phone", "")
            except Exception:
                pass

        # Segment sections for focused stage prompts with expanded non-ATS patterns
        section_regex = r'(?im)^\s*[\#\*\-]*\s*(PROFESSIONAL\s+SUMMARY|SUMMARY|ABOUT\s+ME|ABOUT|WORK\s+EXPERIENCE|EXPERIENCE|EMPLOYMENT|INTERNSHIPS|WORK\s+HISTORY|CAREER|PAST\s+ROLES|WHERE\s+I[\'\’]?VE\s+WORKED|EDUCATION|ACADEMIC\s+BACKGROUND|ACADEMICS|STUDIED|WHERE\s+I\s+STUDIED|COLLEGE|UNIVERSITY|TECHNICAL\s+SKILLS|SKILLS|WHAT\s+I\s+DO|TECHNOLOGIES|TOOLS|TECHNICAL\s+PROJECTS|PROJECTS|THINGS\s+I[\'\’]?VE\s+BUILT|PORTFOLIO|CERTIFICATIONS|LICENSES|ACHIEVEMENTS|AWARDS|HONORS|ACCOMPLISHMENTS)\s*:?\s*$'
        matches = list(re.finditer(section_regex, clean_text))

        sections = {}
        for i, m in enumerate(matches):
            header = m.group(1).upper()
            start_pos = m.end()
            end_pos = matches[i+1].start() if i + 1 < len(matches) else len(clean_text)
            sections[header] = clean_text[start_pos:end_pos].strip()

        def get_sec(*keys):
            for k in keys:
                for sec_header, sec_body in sections.items():
                    if k in sec_header:
                        return sec_body
            return ""

        # Stage 2: Skills Extraction
        skills_text = get_sec("SKILLS", "TECHNOLOGIES", "TOOLS")
        if skills_text:
            stage2_prompt = f"""Extract all programming languages, tools and technologies as a JSON array of skill strings:
Text:
\"\"\"
{skills_text[:1000]}
\"\"\"
Return ONLY valid JSON matching this format:
{{"skills": ["Python", "FastAPI", "React", "Docker"]}}"""
            s2_res = run_stage_prompt("Stage 2 - Skills", stage2_prompt)
            if s2_res:
                try:
                    data = json.loads(s2_res)
                    if isinstance(data, dict) and isinstance(data.get("skills"), list):
                        result["skills"] = data["skills"]
                except Exception:
                    pass

        # Stage 3: Work Experience Extraction (Chunked)
        work_text = get_sec("WORK", "EXPERIENCE", "EMPLOYMENT", "INTERNSHIPS", "HISTORY", "CAREER", "ROLES")
        if work_text:
            stage3_prompt = f"""Extract work experience entries as a JSON array of entry blocks:
Text:
\"\"\"
{work_text[:2000]}
\"\"\"
Return ONLY valid JSON matching this format:
{{"work_history": ["Role Title at Company Name (Time Period)\\n• Bullet point 1\\n• Bullet point 2"]}}"""
            s3_res = run_stage_prompt("Stage 3 - Work", stage3_prompt)
            if s3_res:
                try:
                    data = json.loads(s3_res)
                    if isinstance(data, dict) and isinstance(data.get("work_history"), list):
                        result["work_history"] = data["work_history"]
                except Exception:
                    pass

        # Stage 4: Education Extraction
        edu_text = get_sec("EDUCATION", "ACADEMIC", "STUDIED", "COLLEGE", "UNIVERSITY")
        if edu_text:
            stage4_prompt = f"""Extract education entries as a JSON array of entry blocks:
Text:
\"\"\"
{edu_text[:1200]}
\"\"\"
Return ONLY valid JSON matching this format:
{{"education": ["Degree / Course, Institution Name (Time Period)\\nGrades / CGPA / Details"]}}"""
            s4_res = run_stage_prompt("Stage 4 - Education", stage4_prompt)
            if s4_res:
                try:
                    data = json.loads(s4_res)
                    if isinstance(data, dict) and isinstance(data.get("education"), list):
                        result["education"] = data["education"]
                except Exception:
                    pass

        # Stage 5: Projects & Achievements Extraction
        proj_text = get_sec("PROJECTS", "BUILT", "PORTFOLIO", "ACHIEVEMENTS", "AWARDS")
        if proj_text:
            stage5_prompt = f"""Extract projects and achievements as a JSON object:
Text:
\"\"\"
{proj_text[:1500]}
\"\"\"
Return ONLY valid JSON matching this format:
{{"projects": ["Project Name: Description and Tech Stack used"], "achievements": ["Achievement or Award 1"]}}"""
            s5_res = run_stage_prompt("Stage 5 - Projects", stage5_prompt)
            if s5_res:
                try:
                    data = json.loads(s5_res)
                    if isinstance(data, dict):
                        if isinstance(data.get("projects"), list):
                            result["projects"] = data["projects"]
                        if isinstance(data.get("achievements"), list):
                            result["achievements"] = data["achievements"]
                except Exception:
                    pass

        # Stage 6: Certifications & Licenses Extraction
        cert_text = get_sec("CERTIFICATIONS", "LICENSES", "CERTIFICATES")
        if cert_text:
            stage6_prompt = f"""Extract certifications, professional licenses, and credentials as a JSON array of strings:
Text:
\"\"\"
{cert_text[:1200]}
\"\"\"
Return ONLY valid JSON matching this format:
{{"certifications": ["AWS Certified Solutions Architect", "PMP - Project Management Professional"]}}"""
            s6_res = run_stage_prompt("Stage 6 - Certifications", stage6_prompt)
            if s6_res:
                try:
                    data = json.loads(s6_res)
                    if isinstance(data, dict) and isinstance(data.get("certifications"), list):
                        result["certifications"] = data["certifications"]
                except Exception:
                    pass

        return result

    def extract_profile(self, text: str) -> Dict[str, Any]:
        """
        Parses resume text into structured components:
        Name, Email, Phone, Skills, Education, Work Experience, Achievements, Projects & Certifications.
        Ensures complete isolation per candidate profile.
        """
        clean_text = self.sanitize_text(text)
        
        # 1. Attempt LLM-assisted extraction if configured
        llm_data = self.llm_parse_resume(clean_text)

        profile = {
            "full_name": llm_data.get("full_name") or self.extract_candidate_name(clean_text),
            "email": llm_data.get("email") or self.extract_candidate_email(clean_text),
            "phone": llm_data.get("phone") or self.extract_candidate_phone(clean_text),
            "skills": {
                "programming_languages": [],
                "frameworks": [],
                "tools_databases": [],
                "design_and_ai": []
            },
            "education": llm_data.get("education") or [],
            "certifications": llm_data.get("certifications") or [],
            "experience_years": 0,
            "projects": llm_data.get("projects") or [],
            "work_history": llm_data.get("work_history") or [],
            "achievements": llm_data.get("achievements") or [],
            "domain_expertise": []
        }

        profile["experience_years"] = self.extract_experience_years(clean_text)

        # Canonical Skill Taxonomy Extraction
        seen_canonicals = set()
        for canonical, category, patterns, case_sens in SKILL_TAXONOMY:
            if canonical in seen_canonicals:
                continue
            matched = False
            flags = 0 if case_sens else re.IGNORECASE
            for pat in patterns:
                if re.search(pat, clean_text, flags):
                    matched = True
                    break
            if matched:
                seen_canonicals.add(canonical)
                profile["skills"][category].append({"name": canonical, "level": "Intermediate"})

        # Add LLM extracted skills if present
        if llm_data.get("skills"):
            for sk in llm_data["skills"]:
                if isinstance(sk, str) and sk not in seen_canonicals:
                    seen_canonicals.add(sk)
                    profile["skills"]["tools_databases"].append({"name": sk, "level": "Intermediate"})

        # Section Segmentation Fallback if static parsing needed
        section_regex = r'(?im)^\s*[\#\*\-]*\s*(PROFESSIONAL\s+SUMMARY|SUMMARY|ABOUT\s+ME|ABOUT|WORK\s+EXPERIENCE|EXPERIENCE|EMPLOYMENT|INTERNSHIPS|WORK\s+HISTORY|CAREER|PAST\s+ROLES|WHERE\s+I[\'\’]?VE\s+WORKED|EDUCATION|ACADEMIC\s+BACKGROUND|ACADEMICS|STUDIED|WHERE\s+I\s+STUDIED|COLLEGE|UNIVERSITY|TECHNICAL\s+SKILLS|SKILLS|WHAT\s+I\s+DO|TECHNOLOGIES|TOOLS|TECHNICAL\s+PROJECTS|PROJECTS|THINGS\s+I[\'\’]?VE\s+BUILT|PORTFOLIO|CERTIFICATIONS|LICENSES|ACHIEVEMENTS|AWARDS|HONORS|ACCOMPLISHMENTS)\s*:?\s*$'
        matches = list(re.finditer(section_regex, clean_text))

        sections = {}
        for i, m in enumerate(matches):
            header = m.group(1).upper()
            start_pos = m.end()
            end_pos = matches[i+1].start() if i + 1 < len(matches) else len(clean_text)
            sections[header] = clean_text[start_pos:end_pos].strip()

        def get_section_content(*keys):
            for k in keys:
                for sec_header, sec_body in sections.items():
                    if k in sec_header:
                        return sec_body
            return ""

    def _split_section_into_entries(self, section_text: str) -> List[str]:
        """
        Splits section text into discrete individual entry strings.
        Supports explicit delimiter '<<<KEIKO_ITEM_BREAK>>>' as well as double newlines and entry titles.
        """
        if not section_text:
            return []

        # Strip section tags before splitting
        clean_text = re.sub(r'<<<KEIKO_SECTION_START:[^>]+>>>', '', section_text, flags=re.IGNORECASE)
        clean_text = re.sub(r'<<<KEIKO_SECTION_END:[^>]+>>>', '', clean_text, flags=re.IGNORECASE).strip()

        if not clean_text:
            return []

        initial_chunks = [clean_text]
        if "<<<KEIKO_ITEM_BREAK>>>" in clean_text:
            initial_chunks = clean_text.split("<<<KEIKO_ITEM_BREAK>>>")

        entries = []

        for chunk in initial_chunks:
            c_str = chunk.strip()
            if not c_str:
                continue

            blocks = re.split(r'\n\s*\n', c_str)
            for b in blocks:
                lines = [l.strip() for l in b.splitlines() if l.strip()]
                if not lines:
                    continue

                current_entry = []
                for line in lines:
                    is_start = False
                    if current_entry:
                        colon_match = re.match(r"^([^:\n]+):\s+(.*)$", line)
                        if colon_match and len(colon_match.group(1)) <= 65 and not line.startswith(("-", "*", "•")):
                            is_start = True
                        elif re.search(r"\b(20\d{2}|19\d{2})\s*[-–—]\s*(Present|Current|20\d{2}|19\d{2})\b", line, re.IGNORECASE):
                            is_start = True
                        elif (re.search(r"\b(at|@|\|)\b", line, re.IGNORECASE) or re.match(r"^[A-Z][A-Za-z0-9\s,\-\.]{3,50}\s*[-–—]\s*[A-Z][A-Za-z0-9\s,\-\.]{3,50}", line)) and not line.startswith(("-", "*", "•")):
                            is_start = True
                        elif re.match(r"^(?:\d+[\.\)]|\[\d+\]|Item\s+\d+|Project\s+\d+)\s+", line, re.IGNORECASE):
                            is_start = True

                    if is_start and current_entry:
                        entries.append("\n".join(current_entry))
                        current_entry = [line]
                    else:
                        current_entry.append(line)

                if current_entry:
                    entries.append("\n".join(current_entry))

        return [e for e in entries if len(e.strip()) > 3]

    def extract_profile(self, text: str) -> Dict[str, Any]:
        """
        Parses resume text into structured components:
        Name, Email, Phone, Skills, Education, Work Experience, Achievements, Projects & Certifications.
        Ensures complete isolation per candidate profile.
        """
        clean_text = self.sanitize_text(text)
        
        # 1. Attempt LLM-assisted extraction if configured
        llm_data = self.llm_parse_resume(clean_text)

        profile = {
            "full_name": llm_data.get("full_name") or self.extract_candidate_name(clean_text),
            "email": llm_data.get("email") or self.extract_candidate_email(clean_text),
            "phone": llm_data.get("phone") or self.extract_candidate_phone(clean_text),
            "skills": {
                "programming_languages": [],
                "frameworks": [],
                "tools_databases": [],
                "design_and_ai": []
            },
            "education": llm_data.get("education") or [],
            "certifications": llm_data.get("certifications") or [],
            "experience_years": 0,
            "projects": llm_data.get("projects") or [],
            "work_history": llm_data.get("work_history") or [],
            "achievements": llm_data.get("achievements") or [],
            "domain_expertise": []
        }

        profile["experience_years"] = self.extract_experience_years(clean_text)

        # Canonical Skill Taxonomy Extraction
        seen_canonicals = set()
        for canonical, category, patterns, case_sens in SKILL_TAXONOMY:
            if canonical in seen_canonicals:
                continue
            matched = False
            flags = 0 if case_sens else re.IGNORECASE
            for pat in patterns:
                if re.search(pat, clean_text, flags):
                    matched = True
                    break
            if matched:
                seen_canonicals.add(canonical)
                profile["skills"][category].append({"name": canonical, "level": "Intermediate"})

        # Add LLM extracted skills if present
        if llm_data.get("skills"):
            for sk in llm_data["skills"]:
                if isinstance(sk, str) and sk not in seen_canonicals:
                    seen_canonicals.add(sk)
                    profile["skills"]["tools_databases"].append({"name": sk, "level": "Intermediate"})

        # Section Segmentation Fallback if static parsing needed
        section_regex = r'(?im)^\s*[\#\*\-]*\s*(PROFESSIONAL\s+SUMMARY|SUMMARY|ABOUT\s+ME|ABOUT|WORK\s+EXPERIENCE|EXPERIENCE|EMPLOYMENT|INTERNSHIPS|WORK\s+HISTORY|CAREER|PAST\s+ROLES|WHERE\s+I[\'\’]?VE\s+WORKED|EDUCATION|ACADEMIC\s+BACKGROUND|ACADEMICS|STUDIED|WHERE\s+I\s+STUDIED|COLLEGE|UNIVERSITY|TECHNICAL\s+SKILLS|SKILLS|WHAT\s+I\s+DO|TECHNOLOGIES|TOOLS|TECHNICAL\s+PROJECTS|PROJECTS|THINGS\s+I[\'\’]?VE\s+BUILT|PORTFOLIO|CERTIFICATIONS|LICENSES|ACHIEVEMENTS|AWARDS|HONORS|ACCOMPLISHMENTS)\s*:?\s*$'
        matches = list(re.finditer(section_regex, clean_text))

        sections = {}
        for i, m in enumerate(matches):
            header = m.group(1).upper()
            start_pos = m.end()
            end_pos = matches[i+1].start() if i + 1 < len(matches) else len(clean_text)
            sections[header] = clean_text[start_pos:end_pos].strip()

        def get_section_content(*keys):
            for k in keys:
                for sec_header, sec_body in sections.items():
                    if k in sec_header:
                        return sec_body
            return ""

        if not profile["education"]:
            edu_body = get_section_content("EDUCATION", "ACADEMIC", "STUDIED", "COLLEGE", "UNIVERSITY")
            if edu_body:
                profile["education"] = self._split_section_into_entries(edu_body)

        if not profile["work_history"]:
            exp_body = get_section_content("WORK", "EXPERIENCE", "EMPLOYMENT", "INTERNSHIPS", "HISTORY", "CAREER", "ROLES")
            if exp_body:
                profile["work_history"] = self._split_section_into_entries(exp_body)

        if not profile["projects"]:
            proj_body = get_section_content("PROJECTS", "BUILT", "PORTFOLIO", "TECHNICAL PROJECTS")
            if proj_body:
                entries = self._split_section_into_entries(proj_body)
                for b in entries:
                    if b.strip():
                        lines = [l.strip() for l in b.splitlines() if l.strip()]
                        p_name = lines[0] if lines else "Project"
                        p_name = re.sub(r"^[\bullet\*\-\#\:\s]+", "", p_name).strip()
                        profile["projects"].append({"name": p_name[:70], "description": b.strip()})

        if not profile["certifications"]:
            cert_body = get_section_content("CERTIFICATIONS", "LICENSES", "CERTIFICATES")
            if cert_body:
                profile["certifications"] = self._split_section_into_entries(cert_body)

        if not profile["achievements"]:
            ach_body = get_section_content("ACHIEVEMENTS", "AWARDS", "HONORS", "ACCOMPLISHMENTS")
            if ach_body:
                profile["achievements"] = self._split_section_into_entries(ach_body)

        # Apply Block Classifier & Flipping Engine
        profile = self._reclassify_and_distribute(profile, clean_text)

        # Determine Domain Expertise dynamically from extracted skills
        extracted_skill_names = [s["name"] for cat in profile["skills"].values() for s in cat]
        if any(s in extracted_skill_names for s in ["Python", "FastAPI", "PostgreSQL", "SQL"]):
            profile["domain_expertise"].append("Backend & API Development")
        if any(s in extracted_skill_names for s in ["React", "HTML/CSS", "Tailwind CSS", "UI/UX Design"]):
            profile["domain_expertise"].append("Frontend & UI/UX Design")
        if any(s in extracted_skill_names for s in ["PyTorch", "OpenCV", "TensorFlow"]):
            profile["domain_expertise"].append("AI & Computer Vision")
        if any(s in extracted_skill_names for s in ["Bluetooth / BLE", "Wi-Fi & BSSID", "PWA Development"]):
            profile["domain_expertise"].append("IoT & Wireless Systems")

        return profile
